"""Tests for document loaders and the registry dispatch."""

from __future__ import annotations

import pytest

from ragkb.loaders import load_directory, load_document


def test_markdown_loader(tmp_path) -> None:
    path = tmp_path / "guide.md"
    path.write_text("# Title\n\nSome body text.", encoding="utf-8")
    documents = load_document(path)
    assert len(documents) == 1
    assert "Some body text" in documents[0].content
    assert documents[0].metadata["format"] == "markdown"


def test_docx_loader(tmp_path) -> None:
    from docx import Document as DocxDocument

    path = tmp_path / "guide.docx"
    document = DocxDocument()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.save(str(path))

    documents = load_document(path)
    assert len(documents) == 1
    assert "First paragraph" in documents[0].content
    assert documents[0].metadata["format"] == "docx"


def test_pdf_loader(tmp_path) -> None:
    import pymupdf

    path = tmp_path / "guide.pdf"
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Clock configuration chapter")
    pdf.save(str(path))
    pdf.close()

    documents = load_document(path)
    assert len(documents) == 1
    assert "Clock configuration" in documents[0].content
    assert documents[0].metadata["page"] == 1


def test_unsupported_extension_raises(tmp_path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("hello", encoding="utf-8")
    with pytest.raises(ValueError, match="Unsupported file type"):
        load_document(path)


def test_missing_file_raises(tmp_path) -> None:
    with pytest.raises(FileNotFoundError):
        load_document(tmp_path / "nope.md")


def test_load_directory_skips_unsupported(tmp_path) -> None:
    (tmp_path / "a.md").write_text("# A\n\ntext a", encoding="utf-8")
    (tmp_path / "b.txt").write_text("ignored", encoding="utf-8")
    documents = load_directory(tmp_path)
    assert len(documents) == 1
    assert "text a" in documents[0].content
