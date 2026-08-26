"""DOCX loader (.docx; legacy .doc needs LibreOffice conversion, later milestone)."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from ragkb.core.models import Document
from ragkb.loaders.base import Loader


class DOCXLoader(Loader):
    """Extracts paragraphs and table text from a .docx file."""

    def load(self, path: Path) -> list[Document]:
        document = DocxDocument(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
        content = "\n\n".join(parts)
        if not content.strip():
            return []
        return [
            Document(
                content=content,
                metadata={"source": str(path), "format": "docx", "title": path.stem},
            )
        ]
